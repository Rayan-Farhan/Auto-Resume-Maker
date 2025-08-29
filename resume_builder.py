import json, re, argparse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def normalize(text: str):
    return re.findall(r"[a-z0-9\+#\.]+", text.lower())

def tokenize(text: str):
    return set(normalize(text))

def score_item(jd_tokens, item, fields, tags_field="tags"):
    score = 0
    tags = [t.lower() for t in item.get(tags_field, [])]
    score += 3 * len(set(tags) & jd_tokens)
    for f in fields:
        content = str(item.get(f, "")).lower()
        content_tokens = set(normalize(content))
        score += len(content_tokens & jd_tokens)
    return score

def select_relevant(kb, jd_text, limits):
    jd_tokens = tokenize(jd_text)

    kb_skills = kb.get("skills", [])
    matched_skills = [s for s in kb_skills if s.lower() in jd_tokens]
    if len(matched_skills) < limits.get("skills", 10):
        for s in kb_skills:
            if s in matched_skills:
                continue
            if set(normalize(s)) & jd_tokens:
                matched_skills.append(s)
    matched_skills = matched_skills[:limits.get("skills", 10)]

    def rank_section(items, fields, limit):
        scored = [(score_item(jd_tokens, it, fields), it) for it in items]
        scored = [x for x in scored if x[0] > 0]
        scored.sort(key=lambda x: x[0], reverse=True)
        return [it for _, it in scored[:limit]]

    projects = rank_section(kb.get("projects", []), ["description", "name", "tech"], limits.get("projects", 3))
    experience = rank_section(kb.get("experience", []), ["description","role","company"], limits.get("experience", 3))
    certs = rank_section(kb.get("certificates", []), ["name","issuer"], limits.get("certificates", 3))
    education = kb.get("education", [])[:limits.get("education", 2)]
    return matched_skills, projects, experience, certs, education, jd_tokens

def make_summary(jd_tokens, skills, projects, experience, kb):
    key_terms = sorted(list(jd_tokens & set([s.lower() for s in skills])))[:6]
    parts = []
    if experience:
        role = experience[0].get("role","Engineer")
        parts.append(f"{role} with hands-on work in " + (", ".join(key_terms) if key_terms else "relevant tools"))
    elif projects:
        parts.append("Engineer with projects in " + (", ".join(key_terms) if key_terms else "relevant areas"))
    if projects:
        parts.append(f"Delivered {len(projects)} relevant project(s).")
    return " ".join(parts)

def set_font(run, name="Calibri", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold

def render_word(kb, selected):
    skills, projects, experience, certs, education, jd_tokens = selected
    contact = kb.get("contact", {})
    name = kb.get("name", "Your Name")
    
    doc = Document()
    
    name_para = doc.add_paragraph()
    run = name_para.add_run(name)
    set_font(run, size=18, bold=True)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact Info
    contact_fields = [
        contact.get("email", ""), 
        contact.get("phone", ""), 
        contact.get("github", ""), 
        contact.get("linkedin", "")
    ]
    contact_line = " | ".join([c for c in contact_fields if c])
    if contact_line:
        contact_para = doc.add_paragraph(contact_line)
        set_font(contact_para.runs[0], size=10)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Summary
    summary = make_summary(jd_tokens, skills, projects, experience, kb)
    if summary:
        heading = doc.add_paragraph("SUMMARY")
        set_font(heading.runs[0], size=12, bold=True)
        para = doc.add_paragraph(summary)
        set_font(para.runs[0], size=11)
        doc.add_paragraph()
    
    # Skills
    if skills:
        heading = doc.add_paragraph("SKILLS")
        set_font(heading.runs[0], size=12, bold=True)
        para = doc.add_paragraph(", ".join(skills))  # inline, not bullets
        set_font(para.runs[0], size=11)
        doc.add_paragraph()
    
    # Experience
    if experience:
        heading = doc.add_paragraph("EXPERIENCE")
        set_font(heading.runs[0], size=12, bold=True)
        for e in experience:
            title_line = f"{e.get('role', '')}, {e.get('company', '')} — {e.get('duration', '')}"
            title = doc.add_paragraph()
            run = title.add_run(title_line)
            set_font(run, size=11, bold=True)
            if e.get("description"):
                para = doc.add_paragraph(e["description"])
                set_font(para.runs[0], size=11)
            for b in e.get("impact", []):
                bullet = doc.add_paragraph(f"- {b}", style="List Bullet")
                set_font(bullet.runs[0], size=11)
        doc.add_paragraph()
    
    # Education
    if education:
        heading = doc.add_paragraph("EDUCATION")
        set_font(heading.runs[0], size=12, bold=True)
        for ed in education:
            degree = ed.get("degree", "")
            uni = ed.get("university", "")
            year = f" ({ed.get('year', '')})" if ed.get("year") else ""
            
            if degree and uni:
                edu_line = f"{degree} — {uni}{year}"
            elif degree:
                edu_line = f"{degree}{year}"
            elif uni:
                edu_line = f"{uni}{year}"
            else:
                edu_line = f"Education{year}"  # fallback
            
            para = doc.add_paragraph()
            run = para.add_run(edu_line)
            set_font(run, size=11)
        doc.add_paragraph()

    # Projects
    if projects:
        heading = doc.add_paragraph("PROJECTS")
        set_font(heading.runs[0], size=12, bold=True)
        for p in projects:
            proj = doc.add_paragraph()
            run = proj.add_run(p.get("name", ""))
            set_font(run, size=11, bold=True)
            if p.get("description"):
                para = doc.add_paragraph(p["description"])
                set_font(para.runs[0], size=11)
            tech = p.get("tech", [])
            if tech:
                para = doc.add_paragraph(f"Tech: {', '.join(tech)}")
                set_font(para.runs[0], size=11)
        doc.add_paragraph()
    
    # Certifications
    if certs:
        heading = doc.add_paragraph("CERTIFICATIONS")
        set_font(heading.runs[0], size=12, bold=True)
        for c in certs:
            name = c.get("name", "")
            issuer = c.get("issuer", "")
            year = f" ({c.get('year', '')})" if c.get("year") else ""
            
            if name and issuer:
                cert_line = f"{name} — {issuer}{year}"
            elif name:
                cert_line = f"{name}{year}"
            elif issuer:
                cert_line = f"{issuer}{year}"
            else:
                cert_line = f"Certificate{year}"  # fallback
            
            para = doc.add_paragraph()
            run = para.add_run(cert_line)
            set_font(run, size=11)
        doc.add_paragraph()
    
    doc.save("resume.docx")
    print("Saved as resume.docx")

def main():
    parser = argparse.ArgumentParser(description="Tailor a resume to a job description using a personal knowledge base.")
    parser.add_argument("--kb", required=True, help="Path to kb.json (same folder as script)")
    parser.add_argument("--jd", required=True, help="Path to job description .txt file (same folder as script)")
    parser.add_argument("--skills", type=int, default=15)
    parser.add_argument("--projects", type=int, default=5)
    parser.add_argument("--experience", type=int, default=4)
    parser.add_argument("--certificates", dest="certs", type=int, default=7)
    parser.add_argument("--education", type=int, default=2)
    args = parser.parse_args()

    with open(args.kb, "r", encoding="utf-8") as f:
        kb = json.load(f)
    with open(args.jd, "r", encoding="utf-8") as f:
        jd_text = f.read()

    limits = {
        "skills": args.skills,
        "projects": args.projects,
        "experience": args.experience,
        "certificates": args.certs,
        "education": args.education,
    }

    selected = select_relevant(kb, jd_text, limits)
    
    render_word(kb, selected)

if __name__ == "__main__":
    main()